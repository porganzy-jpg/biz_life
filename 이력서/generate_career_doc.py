"""
경력기술서 (Career Description) Generator — v2 (Polished)
연새한솔 | Senior Data Analyst & AI Strategy Manager
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ─── Constants ───
BLUE_ACCENT = RGBColor(0x25, 0x63, 0xEB)
DARK_BLUE = RGBColor(0x1E, 0x40, 0xAF)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MEDIUM_GRAY = RGBColor(0x55, 0x55, 0x55)
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
SUCCESS_GREEN = RGBColor(0x05, 0x96, 0x69)
HIGHLIGHT_ORANGE = RGBColor(0xEA, 0x58, 0x0C)
TABLE_HEADER_BG = "2563EB"
TABLE_ALT_BG = "F0F4FF"
TABLE_BORDER_COLOR = "B0C4DE"
HIGHLIGHT_BG = "EFF6FF"
RESULT_BG = "F0FDF4"
FORMULA_BG = "FFF7ED"
FONT_KR = "맑은 고딕"
FONT_EN = "Calibri"

IMG_DIR = r"C:/Users/user/Desktop/이력서/images"

doc = Document()

# ─── Page Setup ───
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ─── Style Definitions ───
style = doc.styles['Normal']
style.font.name = FONT_EN
style.font.size = Pt(10.5)
style.font.color.rgb = DARK_GRAY
style.paragraph_format.line_spacing = 1.4
style.paragraph_format.space_after = Pt(4)
rFonts = style.element.rPr.rFonts if style.element.rPr is not None else style.element.get_or_add_rPr().get_or_add_rFonts()
rFonts.set(qn('w:eastAsia'), FONT_KR)

h1 = doc.styles['Heading 1']
h1.font.name = FONT_EN; h1.font.size = Pt(18); h1.font.bold = True; h1.font.color.rgb = BLUE_ACCENT
h1.paragraph_format.space_before = Pt(24); h1.paragraph_format.space_after = Pt(12)
h1.paragraph_format.keep_with_next = True
h1.element.rPr.get_or_add_rFonts().set(qn('w:eastAsia'), FONT_KR)

h2 = doc.styles['Heading 2']
h2.font.name = FONT_EN; h2.font.size = Pt(14); h2.font.bold = True; h2.font.color.rgb = BLUE_ACCENT
h2.paragraph_format.space_before = Pt(18); h2.paragraph_format.space_after = Pt(8)
h2.paragraph_format.keep_with_next = True
h2.element.rPr.get_or_add_rFonts().set(qn('w:eastAsia'), FONT_KR)

h3 = doc.styles['Heading 3']
h3.font.name = FONT_EN; h3.font.size = Pt(12); h3.font.bold = True; h3.font.color.rgb = DARK_BLUE
h3.paragraph_format.space_before = Pt(14); h3.paragraph_format.space_after = Pt(6)
h3.paragraph_format.keep_with_next = True
h3.element.rPr.get_or_add_rFonts().set(qn('w:eastAsia'), FONT_KR)


# ════════════════════════════════════════════════════════════════
#  HELPER FUNCTIONS
# ════════════════════════════════════════════════════════════════

def _set_font_kr(run):
    rpr = run._element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn('w:eastAsia'), FONT_KR)


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def format_table_borders(table, color="B0C4DE"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def format_cell_text(cell, text, bold=False, size=Pt(10), color=DARK_GRAY, alignment=None, font_name=FONT_EN):
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment: p.alignment = alignment
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = font_name; run.font.size = size; run.font.bold = bold; run.font.color.rgb = color
    _set_font_kr(run)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def format_cell_rich(cell, parts, alignment=None):
    """Format cell with mixed styling. parts = [(text, bold, size, color), ...]"""
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment: p.alignment = alignment
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    for text, bold, size, color in parts:
        run = p.add_run(text)
        run.font.name = FONT_EN; run.font.size = size; run.font.bold = bold; run.font.color.rgb = color
        _set_font_kr(run)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_blue_line(doc_obj):
    p = doc_obj.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(8)
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="1" w:color="2563EB"/></w:pBdr>')
    p._element.get_or_add_pPr().append(pBdr)
    return p


def add_thin_gray_line(doc_obj):
    p = doc_obj.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="D1D5DB"/></w:pBdr>')
    p._element.get_or_add_pPr().append(pBdr)
    return p


def add_styled_paragraph(doc_obj, text, size=Pt(10.5), bold=False, color=DARK_GRAY,
                         alignment=None, space_before=Pt(0), space_after=Pt(4),
                         line_spacing=1.4, first_line_indent=None):
    p = doc_obj.add_paragraph()
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = line_spacing
    if alignment: p.alignment = alignment
    if first_line_indent: p.paragraph_format.first_line_indent = first_line_indent
    run = p.add_run(text)
    run.font.name = FONT_EN; run.font.size = size; run.font.bold = bold; run.font.color.rgb = color
    _set_font_kr(run)
    return p


def add_mixed_paragraph(doc_obj, parts, space_before=Pt(0), space_after=Pt(4),
                        line_spacing=1.4, alignment=None, left_indent=None):
    p = doc_obj.add_paragraph()
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = line_spacing
    if alignment: p.alignment = alignment
    if left_indent: p.paragraph_format.left_indent = left_indent
    for text, bold, size, color in parts:
        run = p.add_run(text)
        run.font.name = FONT_EN; run.font.size = size; run.font.bold = bold; run.font.color.rgb = color
        _set_font_kr(run)
    return p


def add_bullet_point(doc_obj, text, level=0, bold_prefix=None, size=Pt(10)):
    indent = Cm(1.0 + level * 0.8)
    p = doc_obj.add_paragraph()
    p.paragraph_format.left_indent = indent
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.35
    bullet_char = "•" if level == 0 else "–" if level == 1 else "▸"
    if bold_prefix:
        r1 = p.add_run(f"{bullet_char} {bold_prefix}")
        r1.font.name = FONT_EN; r1.font.size = size; r1.font.bold = True; r1.font.color.rgb = DARK_GRAY
        _set_font_kr(r1)
        r2 = p.add_run(text)
        r2.font.name = FONT_EN; r2.font.size = size; r2.font.color.rgb = DARK_GRAY
        _set_font_kr(r2)
    else:
        r = p.add_run(f"{bullet_char} {text}")
        r.font.name = FONT_EN; r.font.size = size; r.font.color.rgb = DARK_GRAY
        _set_font_kr(r)
    return p


def add_project_image(doc_obj, image_path, caption, width=Inches(5.5)):
    if not os.path.exists(image_path):
        return
    p_img = doc_obj.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(8)
    p_img.paragraph_format.space_after = Pt(2)
    run = p_img.add_run()
    run.add_picture(image_path, width=width)
    p_cap = doc_obj.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(2)
    p_cap.paragraph_format.space_after = Pt(10)
    cap_run = p_cap.add_run(caption)
    cap_run.font.name = FONT_EN; cap_run.font.size = Pt(9); cap_run.font.italic = True; cap_run.font.color.rgb = LIGHT_GRAY
    _set_font_kr(cap_run)


# ─── NEW: Enhanced formatting helpers ───

def add_highlight_title(doc_obj, number, title):
    """큰 하이라이트 제목 (★ Highlight 1: ...) — 파란 배경 띠"""
    table = doc_obj.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "1E40AF")
    format_cell_text(cell, f"  ★  Highlight {number}: {title}", bold=True, size=Pt(12),
                     color=WHITE, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    # remove table borders for clean look
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    doc_obj.add_paragraph().paragraph_format.space_after = Pt(2)


def add_shaded_box(doc_obj, label, text, bg_color=HIGHLIGHT_BG):
    """라벨 + 본문이 있는 음영 박스 (1행 1열 테이블)"""
    table = doc_obj.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, bg_color)
    format_table_borders(table, "D0D5DD")
    parts = [(f"{label}  ", True, Pt(10), BLUE_ACCENT),
             (text, False, Pt(10), DARK_GRAY)]
    format_cell_rich(cell, parts)
    doc_obj.add_paragraph().paragraph_format.space_after = Pt(1)


def add_result_box(doc_obj, items):
    """성과/임팩트 강조 박스 (녹색 배경)"""
    table = doc_obj.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, RESULT_BG)
    format_table_borders(table, "86EFAC")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.4
    label_run = p.add_run("▶ 핵심 성과   ")
    label_run.font.name = FONT_EN; label_run.font.size = Pt(10); label_run.font.bold = True
    label_run.font.color.rgb = SUCCESS_GREEN
    _set_font_kr(label_run)
    for i, item in enumerate(items):
        if i > 0:
            br_run = p.add_run("\n")
        r = p.add_run(f"• {item}")
        r.font.name = FONT_EN; r.font.size = Pt(9.5); r.font.color.rgb = DARK_GRAY
        _set_font_kr(r)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc_obj.add_paragraph().paragraph_format.space_after = Pt(2)


def add_formula_box(doc_obj, formula_lines):
    """포뮬러/공식 강조 박스 (주황 배경)"""
    table = doc_obj.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, FORMULA_BG)
    format_table_borders(table, "FDBA74")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    for i, line in enumerate(formula_lines):
        if i > 0:
            p.add_run("\n")
        bold = (i == 0)  # first line is the main formula
        r = p.add_run(line)
        r.font.name = FONT_EN; r.font.size = Pt(10) if bold else Pt(9.5)
        r.font.bold = bold; r.font.color.rgb = HIGHLIGHT_ORANGE if bold else DARK_GRAY
        _set_font_kr(r)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc_obj.add_paragraph().paragraph_format.space_after = Pt(2)


def add_step_section(doc_obj, step_num, title, bullets):
    """번호가 매겨진 스텝 섹션 (Step 1. ...)"""
    p = doc_obj.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    r1 = p.add_run(f"Step {step_num}.  ")
    r1.font.name = FONT_EN; r1.font.size = Pt(11); r1.font.bold = True; r1.font.color.rgb = BLUE_ACCENT
    _set_font_kr(r1)
    r2 = p.add_run(title)
    r2.font.name = FONT_EN; r2.font.size = Pt(11); r2.font.bold = True; r2.font.color.rgb = DARK_GRAY
    _set_font_kr(r2)
    for b in bullets:
        add_bullet_point(doc_obj, b, level=1, size=Pt(9.5))


def add_sub_project(doc_obj, title, bullets, tools=None, result=None):
    """서브 프로젝트 (2-1. GMM 같은 하위 항목)"""
    p = doc_obj.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run(f"▎ {title}")
    r.font.name = FONT_EN; r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = DARK_BLUE
    _set_font_kr(r)
    for b in bullets:
        add_bullet_point(doc_obj, b, level=1, size=Pt(9.5))
    if tools:
        add_mixed_paragraph(doc_obj, [
            ("    도구: ", True, Pt(9.5), BLUE_ACCENT),
            (tools, False, Pt(9.5), MEDIUM_GRAY),
        ], space_before=Pt(2), space_after=Pt(2), left_indent=Cm(0.5))
    if result:
        add_mixed_paragraph(doc_obj, [
            ("    → 성과: ", True, Pt(9.5), SUCCESS_GREEN),
            (result, False, Pt(9.5), DARK_GRAY),
        ], space_before=Pt(2), space_after=Pt(2), left_indent=Cm(0.5))


def add_company_highlights(doc_obj, items):
    """회사 헤더 아래 핵심 성과 요약 테이블 (2열: 아이콘+항목)"""
    table = doc_obj.add_table(rows=len(items), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    format_table_borders(table, "E2E8F0")
    for i, (icon, text) in enumerate(items):
        c0 = table.rows[i].cells[0]; c0.width = Cm(1.5)
        c1 = table.rows[i].cells[1]; c1.width = Cm(15.0)
        set_cell_shading(c0, HIGHLIGHT_BG)
        format_cell_text(c0, icon, bold=True, size=Pt(9), color=BLUE_ACCENT,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER)
        if i % 2 == 1:
            set_cell_shading(c1, "FAFBFF")
        format_cell_text(c1, text, bold=False, size=Pt(9.5), color=DARK_GRAY)
    doc_obj.add_paragraph().paragraph_format.space_after = Pt(4)


def add_project_section(doc_obj, title, subsections):
    """기존 프로젝트 섹션 (Coupang, eBay, E-Land용)"""
    p_title = doc_obj.add_paragraph()
    p_title.paragraph_format.space_before = Pt(10)
    p_title.paragraph_format.space_after = Pt(6)
    p_title.paragraph_format.left_indent = Cm(0.3)
    run = p_title.add_run(f"▎ {title}")
    run.font.name = FONT_EN; run.font.size = Pt(11.5); run.font.bold = True; run.font.color.rgb = DARK_BLUE
    _set_font_kr(run)
    for key, value in subsections.items():
        if isinstance(value, str):
            add_mixed_paragraph(doc_obj, [
                (f"    {key}: ", True, Pt(10), BLUE_ACCENT),
                (value, False, Pt(10), DARK_GRAY)
            ], space_before=Pt(2), space_after=Pt(2), left_indent=Cm(0.5))
        elif isinstance(value, list):
            add_mixed_paragraph(doc_obj, [
                (f"    {key}:", True, Pt(10), BLUE_ACCENT)
            ], space_before=Pt(2), space_after=Pt(1), left_indent=Cm(0.5))
            for item in value:
                add_bullet_point(doc_obj, item, level=1, size=Pt(10))


# ════════════════════════════════════════════════════════════════
#  SECTION 1: HEADER
# ════════════════════════════════════════════════════════════════

add_styled_paragraph(doc, "경력기술서", size=Pt(28), bold=True, color=BLUE_ACCENT,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(40), space_after=Pt(4))
add_styled_paragraph(doc, "Career Description & Portfolio", size=Pt(14), bold=False,
                    color=LIGHT_GRAY, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    space_before=Pt(0), space_after=Pt(16))
add_blue_line(doc)

add_styled_paragraph(doc, "연새한솔", size=Pt(24), bold=True, color=DARK_GRAY,
                    space_before=Pt(16), space_after=Pt(4))
add_styled_paragraph(doc, "Senior Data Analyst & AI Strategy Manager", size=Pt(13),
                    bold=False, color=BLUE_ACCENT, space_before=Pt(0), space_after=Pt(12))

add_mixed_paragraph(doc, [
    ("010-6821-1779", False, Pt(10.5), MEDIUM_GRAY),
    ("  |  ", False, Pt(10.5), LIGHT_GRAY),
    ("itziana@naver.com", False, Pt(10.5), MEDIUM_GRAY),
], space_after=Pt(4))

add_mixed_paragraph(doc, [
    ("15년+ 데이터 전문가", True, Pt(10.5), DARK_GRAY),
    ("  |  ", False, Pt(10.5), LIGHT_GRAY),
    ("E-Land → eBay Korea → Coupang → Tving", False, Pt(10.5), MEDIUM_GRAY),
], space_after=Pt(20))

add_thin_gray_line(doc)


# ════════════════════════════════════════════════════════════════
#  SECTION 2: PROFESSIONAL SUMMARY
# ════════════════════════════════════════════════════════════════

doc.add_heading("Professional Summary", level=1)

summary_p1 = (
    "오프라인 리테일(E-Land) → e-Commerce(eBay Korea) → 로켓커머스(Coupang) → OTT(Tving)까지, "
    "15년간 산업의 디지털 전환 최전방에서 데이터 전략을 설계하고 실행해 온 분석 리더입니다. "
    "각 회사에서 '데이터로 풀어야 할 가장 중요한 문제'를 정의하고, 분석 체계를 직접 구축하며, "
    "그 결과를 비즈니스 성과로 연결하는 전 과정을 일관되게 수행해 왔습니다."
)
summary_p2 = (
    "커리어 전반에 걸쳐 반복되는 패턴은 '시스템적 문제 해결'입니다. "
    "eBay Korea에서는 3개 사이트의 데이터 거버넌스가 통일되지 않아 전사 의사결정이 불가능한 상황을 "
    "행동데이터 수집 시스템 구축과 Standard Query 체계 정립으로 해결했고, "
    "Coupang에서는 3P Marketplace의 셀러 성장 정체 문제를 AARRR 퍼널 분석과 AB Test를 통해 "
    "구조적 병목을 찾아내어 해소했습니다. 기술(SQL, Python, Tableau, Airflow)을 도구로 쓰되, "
    "항상 비즈니스 질문에서 출발하여 실행 가능한 전략으로 귀결시키는 것이 저의 분석 철학입니다."
)
summary_p3 = (
    "현재 Tving에서는 이 경험을 AI로 확장하고 있습니다. Tableau 정적 대시보드를 AI가 KPI 이상을 "
    "자동 감지·분석하는 지능형 시스템으로 전환했고, 구독자·콘텐츠 실적의 Bottom-Up 예측 모델을 "
    "구축하여 경영진 의사결정의 정량적 근거를 제공하고 있습니다. 나아가 전사 KPI를 조직별 책임 지표로 "
    "분해하는 OKR 포뮬러를 설계하여, '누가 무엇을 개선해야 하는가'를 데이터가 직접 답하는 체계를 "
    "만들었습니다."
)

add_styled_paragraph(doc, summary_p1, size=Pt(10.5), line_spacing=1.5,
                    space_before=Pt(8), space_after=Pt(10))
add_styled_paragraph(doc, summary_p2, size=Pt(10.5), line_spacing=1.5,
                    space_before=Pt(0), space_after=Pt(10))
add_styled_paragraph(doc, summary_p3, size=Pt(10.5), line_spacing=1.5,
                    space_before=Pt(0), space_after=Pt(16))


# ════════════════════════════════════════════════════════════════
#  SECTION 3: CAREER SUMMARY TABLE
# ════════════════════════════════════════════════════════════════

doc.add_heading("경력 요약", level=1)
add_styled_paragraph(doc, "총 경력 15년+ | 리테일 → e-Commerce → OTT | 데이터 분석 → AI 전략 리더십",
                    size=Pt(10), color=MEDIUM_GRAY, space_after=Pt(10))

career_data = [
    ["회사", "부서", "직급", "기간"],
    ["Tving", "Data Analyst Team", "팀장", "2025.07 ~ 현재"],
    ["Coupang", "Business Intelligence", "Principal", "2020.03 ~ 2025.07"],
    ["eBay Korea", "Data & Martech", "과장 (Data PM)", "2014.07 ~ 2020.03"],
    ["E-Land Retail", "Marketing / CRM", "주임", "2010.11 ~ 2014.07"],
]

table = doc.add_table(rows=5, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
format_table_borders(table, TABLE_BORDER_COLOR)
for row in table.rows:
    row.cells[0].width = Cm(3.5); row.cells[1].width = Cm(5.0)
    row.cells[2].width = Cm(3.5); row.cells[3].width = Cm(4.5)
for j, header_text in enumerate(career_data[0]):
    cell = table.rows[0].cells[j]
    set_cell_shading(cell, TABLE_HEADER_BG)
    format_cell_text(cell, header_text, bold=True, size=Pt(10), color=WHITE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
for i, row_data in enumerate(career_data[1:], 1):
    for j, cell_text in enumerate(row_data):
        cell = table.rows[i].cells[j]
        if i % 2 == 0: set_cell_shading(cell, TABLE_ALT_BG)
        alignment = WD_ALIGN_PARAGRAPH.CENTER if j >= 2 else WD_ALIGN_PARAGRAPH.LEFT
        format_cell_text(cell, cell_text, bold=(j==0), size=Pt(10), color=DARK_GRAY, alignment=alignment)


# ════════════════════════════════════════════════════════════════
#  SECTION 4: CORE COMPETENCIES
# ════════════════════════════════════════════════════════════════

doc.add_heading("Core Competencies", level=1)

comp_table = doc.add_table(rows=4, cols=2)
comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
format_table_borders(comp_table, TABLE_BORDER_COLOR)

comp_data = [
    ("Technical Skills", "SQL (Hive/Oracle/MySQL), Tableau, Python, Google Analytics, "
     "Apache Airflow, ElasticSearch, Kibana, R, D3.js, Bootstrap, Zeplin"),
    ("Analysis Methods", "AB Testing, Cohort Analysis, AARRR Framework, Funnel Analysis, "
     "Segmentation, Fraud Detection, CRM Analytics, Data Governance, KPI Design"),
    ("AI & Automation", "AI 자동화 대시보드, NL→SQL 자동 쿼리, LLM 챗봇, GMM 클러스터링, "
     "예측 모델링(Bottom-Up), AI 자동 재학습 파이프라인, KPI 포뮬러화(OKR)"),
    ("Leadership", "Team Managing, Silo Process Design, Cross-functional Collaboration, "
     "전사 데이터 전략 수립, 교육 및 프레젠테이션"),
]
for i, (label, desc) in enumerate(comp_data):
    cat_cell = comp_table.rows[i].cells[0]; cat_cell.width = Cm(3.5)
    set_cell_shading(cat_cell, "EFF6FF")
    format_cell_text(cat_cell, label, bold=True, size=Pt(10), color=BLUE_ACCENT, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    desc_cell = comp_table.rows[i].cells[1]; desc_cell.width = Cm(13.0)
    if i % 2 == 1: set_cell_shading(desc_cell, "FAFBFF")
    format_cell_text(desc_cell, desc, bold=False, size=Pt(10), color=DARK_GRAY)


# ════════════════════════════════════════════════════════════════
#  PAGE BREAK → SECTION 5: TVING
# ════════════════════════════════════════════════════════════════
doc.add_page_break()

doc.add_heading("경력 상세", level=1)
add_blue_line(doc)

# --- Tving Header ---
doc.add_heading("1. Tving (2025.07 ~ 현재)", level=2)

tving_info = [
    ("직급/역할: ", "Data Analyst Team 팀장 (Manager)"),
    ("조직: ", "Tving Data Analyst Team"),
    ("핵심 역할: ", "AI 기반 지능형 분석 체계 구축, 예측 모델링, KPI 전략 설계, 팀 매니징"),
]
for label, value in tving_info:
    add_mixed_paragraph(doc, [
        (label, True, Pt(10), BLUE_ACCENT),
        (value, False, Pt(10), DARK_GRAY),
    ], space_before=Pt(2), space_after=Pt(2))

# 핵심 성과 요약
add_company_highlights(doc, [
    ("AI Dashboard", "Tableau 정적 대시보드 → AI 지능형 대시보드 전환 (자동 인사이트 생성 + AI 챗봇)"),
    ("Auto Query", "자연어 → SQL 자동 생성·실행하는 AI 자동화 쿼리 시스템 구축"),
    ("Prediction", "구독자·콘텐츠 실적 Bottom-Up 예측 모델 구축, 예측 커버율 ~75% 달성"),
    ("Clustering", "GMM 기반 고객 클러스터링 → 클러스터별 LTV 산출, 이탈 위험군 조기 감지"),
    ("OKR Formula", "전사 KPI를 조직별 책임 지표로 분해하는 OKR 포뮬러 설계·적용"),
    ("Team Lead", "사일로형 업무 프로세스 도입으로 팀 생산성과 전문성 동시 향상"),
])

add_thin_gray_line(doc)

# ── Highlight 1: AI 대시보드 자동화 ──
add_highlight_title(doc, 1, "AI 기반 자동화 대시보드 — 인사이트 자동 생성 & AI 챗봇")

add_shaded_box(doc, "배경 & 문제:",
    "OTT 시장은 구독자 이탈률이 월 5~10%에 달하는 고회전 비즈니스입니다. KPI 이상 신호를 하루라도 "
    "늦게 포착하면 대응 시점을 놓치게 됩니다. 그러나 기존 Tableau 대시보드는 '지표가 떨어졌다'는 사실만 "
    "보여줄 뿐, '왜 떨어졌고, 어떤 팀이 무엇을 바꿔야 하는가'까지 도달하려면 분석가의 수동 탐색에 "
    "수일이 소요되는 구조적 한계가 있었습니다.")

add_shaded_box(doc, "목표:",
    "분석가가 직접 탐색하지 않아도, AI가 KPI 이상을 감지하는 순간 원인을 분석하고 "
    "개선 방향까지 제시하는 '능동형 인사이트 시스템' 구축 — 대응 Lead Time을 수일에서 실시간으로 단축")

add_step_section(doc, 1, "AI 자동화 쿼리 시스템 — 모든 AI 기능의 기반 인프라 구축", [
    "자연어 질문 → SQL 자동 생성·실행 체계를 설계하여, 분석가가 아닌 누구라도 데이터에 접근할 수 있는 계층을 구축",
    "전사 데이터 카탈로그(테이블 스키마, 컬럼 정의, 관계 매핑)를 구조화하여 AI 컨텍스트로 주입 — 쿼리 정확도와 일관성 확보",
    "eBay Korea Montelena TF에서 3개 사이트 데이터 거버넌스를 통일하고 Standard Query를 정립한 경험이 이 설계의 직접적 기반",
    "이 쿼리 시스템이 이후 모든 AI 기능(대시보드 자동 생성, AI 챗봇, 원인 분석)의 공통 데이터 접근 계층으로 작동",
])

add_step_section(doc, 2, "KPI 자동 모니터링 & 원인 분석 — 수동 탐색을 제거", [
    "구독자 수, MAU, 시청시간, 재구독률, 콘텐츠별 실적 등 핵심 KPI를 AI가 상시 트래킹",
    "이상 변동 감지 시 3단계 자동 분석 수행: ① 시계열 분해(Trend/Seasonality/Residual), "
     "② 세그먼트별 기여도 분해(Contribution Analysis), ③ 상관 지표 연관 분석",
    "분석 결과를 '핵심 원인 Top-3 + 차트 + 요약 텍스트'로 자동 구성하여, 의사결정자가 별도 요청 없이 인사이트를 즉시 확인",
])

add_shaded_box(doc, "예시:",
    "'이번 주 재구독률 3%p 하락' → AI 자동 분석 → '25-34세 그룹의 시청시간 15% 감소, "
    "해당 그룹이 주로 시청하던 드라마 A 종영이 주요 원인. 유사 장르 B의 프로모션 노출 강화 권장'",
    bg_color="F5F3FF")

add_step_section(doc, 3, "OKR 목표 대비 진단 — '누가 무엇을 바꿔야 하는가'까지 안내", [
    "전사 OKR 목표 대비 달성도를 실시간 추적하고, 미달 항목에 대해 AI가 구체적 개선 방향을 제시",
    "마케팅·프로덕트·콘텐츠팀 각각의 책임 KPI 중 '어떤 지표가, 얼마나, 왜 부진한지'를 자동 진단",
    "Coupang에서 셀러 퍼널의 이탈 최다 단계를 데이터로 특정하고 개선안을 도출한 방법론을 OTT 구독 맥락으로 확장 적용",
])

add_step_section(doc, 4, "AI 챗봇 — 데이터 민주화의 마지막 퍼즐", [
    "대시보드에 LLM 기반 챗봇을 통합하여, 비분석가도 '지난 달 20대 여성의 시청시간 트렌드는?' 같은 자연어 질문으로 분석 결과를 즉시 확인",
    "지표 정의 확인, 심층 드릴다운, 기간·세그먼트 비교 등 다양한 분석 시나리오를 대화형으로 지원",
    "Step 1의 자동화 쿼리 시스템이 백엔드에서 작동 — 질문 → SQL 생성 → 실행 → 해석 → 시각화를 End-to-End로 자동 처리",
])

add_mixed_paragraph(doc, [
    ("사용 도구: ", True, Pt(9.5), BLUE_ACCENT),
    ("Python, LLM (GPT/Claude API), SQL (자동 생성), Tableau (기존 연동), 자체 대시보드 프레임워크", False, Pt(9.5), MEDIUM_GRAY),
], space_before=Pt(6), space_after=Pt(4), left_indent=Cm(0.5))

add_result_box(doc, [
    "KPI 이상 감지 → 원인 분석 → 개선 방향 제시까지의 Lead Time을 수일 → 실시간으로 단축",
    "분석가의 반복적 탐색 업무를 AI가 대체하여, 분석팀이 전략적 심층 분석에 집중할 수 있는 환경 조성",
    "경영진·PM·마케터가 자연어 챗봇으로 데이터에 직접 접근 가능 — 조직 전체의 데이터 리터러시 향상",
])

add_thin_gray_line(doc)

# ── Highlight 2: 예측 모델링 & KPI 포뮬러화 ──
add_highlight_title(doc, 2, "예측 모델링 체계 구축 & KPI 포뮬러화(OKR) 설계")

add_shaded_box(doc, "배경:",
    "OTT 사업의 핵심 질문은 '다음 분기 구독자가 몇 명이고, 어떤 콘텐츠에 투자해야 하는가'입니다. "
    "그러나 기존에는 단일 시계열 모델이나 경험적 추정에 의존하여 예측 정밀도가 낮았고, "
    "'구독자가 줄었다'는 결과만 확인할 뿐 어떤 조직이 무엇을 개선해야 하는지 분리되지 않는 "
    "구조적 문제가 있었습니다.")

# 2-1. GMM
add_sub_project(doc, "2-1. GMM 기반 고객 클러스터링 & LTV 분석", [
    "Gaussian Mixture Model(GMM)을 활용하여 전체 구독자를 다차원 행동 특성 기반으로 클러스터링",
    "시청 패턴(장르 선호, Binge-watching 여부), 구독 이력(신규/재구독/장기), 서비스 행태(검색 빈도, 디바이스 분포) 등 복합 Feature 설계",
    "클러스터별 프로파일링을 통해 '누가 우리의 핵심 고객이고, 누가 이탈 위험에 있는가'를 정량적으로 정의",
    "클러스터별 LTV 산출 → 고LTV 고객의 공통 행동 패턴 식별, 이탈 위험군에 대한 조기 감지 체계 구축",
    "E-Land에서의 CRM 세그먼트 전략과 Coupang 셀러 Cohort 분석의 경험이 OTT 구독 모델 설계의 기반",
],
    tools="Python (scikit-learn, GMM), SQL",
    result="마케팅 타겟팅 정밀도 향상, 고LTV 고객 유지 전략 수립 및 이탈 위험군 선제 대응 체계 확립")

# 2-2. 예측 모델
add_sub_project(doc, "2-2. 콘텐츠 실적 & 구독자 수 예측 모델 — 커버율 ~75% 달성", [
    "핵심 아이디어: 전체를 한 번에 예측하는 대신, 구성 요소로 분해한 후 각각을 독립 예측하여 합산(Bottom-Up)",
    "구독자 예측 = 신규유입(채널별) + 재구독(세그먼트별) - 이탈(위험군별), 각 항목을 디바이스·연령대로 다시 세분화",
    "콘텐츠 실적도 장르·요일·시간대·독점 여부로 세분화하여, 각 그룹 특성에 맞는 최적 모델을 개별 적용",
    "세분화를 통해 집계 수준의 노이즈를 제거하여, 예측 커버율 ~75%를 달성",
    "AI 자동 재학습(Auto-Retraining) 파이프라인을 구축하여, 시장 변화에 모델이 자동 적응하는 운영 체계 완성",
],
    tools="Python (scikit-learn, statsmodels, Prophet), SQL, AI Auto-Retraining 파이프라인",
    result="경영진 분기 사업계획 수립 시 핵심 정량 근거로 활용 — '감'이 아닌 '데이터'에 기반한 의사결정 지원")

# 2-3. KPI 포뮬러화
add_sub_project(doc, "2-3. KPI 포뮬러화 & OKR 체계 설계 — '누구의 책임인가'를 데이터로 답하다", [
    "Tving 최상위 KPI '구독자 수'의 70% 이상이 재구독자 — 그러나 이 단일 지표가 전사 목표로만 존재하고, 조직별 책임이 불분명",
    "구독자가 줄어도 '마케팅 문제인지, 콘텐츠 문제인지, UX 문제인지' 특정할 수 없어 대응이 지연되는 구조",
    "데이터 분석을 통해 '시청시간'이 재구독 전환의 가장 강력한 선행 지표임을 통계적으로 검증",
    "이 발견을 바탕으로 재구독 KPI를 3개 조직의 책임 지표로 분해하는 포뮬러를 설계",
])

add_formula_box(doc, [
    "재구독자 수  =  ① 재방문 인원 수  ×  ② 시청 전환율  ×  ③ 시청 전환 당 시청시간",
    "① 재방문 인원 수  →  마케팅팀 대표 KPI (CRM, 리타겟팅, 푸시 알림 등)",
    "② 시청 전환율  →  프로덕트팀 대표 KPI (UI/UX, 추천 알고리즘, 온보딩 등)",
    "③ 시청 전환 당 시청시간  →  컨텐츠팀 대표 KPI (컨텐츠 품질, 편성, 독점 투자 등)",
])

add_bullet_point(doc, "각 팀이 '내가 무엇을 개선해야 전사 구독자 수가 올라가는지' 명확히 인지하고 집중할 수 있는 체계 확립", level=1, size=Pt(9.5))
add_bullet_point(doc, "예: '재방문은 유지되나 시청 전환율 하락' → 프로덕트팀이 UI/UX 원인을 파악·개선해야 하는 상황임을 즉시 특정", level=1, size=Pt(9.5))
add_bullet_point(doc, "Coupang 셀러 퍼널 'SignUp → Live → Active → GMV' 단계별 분해 방법론을 구독 비즈니스에 맞게 재설계한 결과", level=1, size=Pt(9.5))

add_result_box(doc, [
    "조직 간 '책임 떠넘기기' 문제 해소 — 데이터에 기반한 명확한 성과 귀인(Attribution) 체계 수립",
    "구독자 수 변동 시 포뮬러로 원인 조직을 즉시 특정 → 해당 팀이 구체적 개선 액션을 자율적으로 도출",
    "Highlight 1의 AI 대시보드와 연동하여, 포뮬러 기반 자동 진단이 실시간으로 작동",
])

add_thin_gray_line(doc)

# ── 팀 매니징 ──
add_sub_project(doc, "데이터 팀 매니징 & 사일로형 프로세스 도입", [
    "팀원별 전문 도메인(콘텐츠, 마케팅, 프로덕트, 커머스 등)을 명확히 매핑하고, 각 영역의 분석 오너십을 부여",
    "사일로 단위 자율 의사결정 권한 위임 — 팀장 승인 병목을 해소하여 분석 속도와 팀원 성장을 동시에 확보",
    "주간 Sync-up(진행 공유) + 월간 Deep Dive(전략적 심층 분석 리뷰) + 팀 내 지식 공유 세션 정례화",
])


# ════════════════════════════════════════════════════════════════
#  PAGE BREAK → COUPANG
# ════════════════════════════════════════════════════════════════
doc.add_page_break()

doc.add_heading("2. Coupang (2020.03 ~ 2025.07)", level=2)

coupang_info = [
    ("직급/역할: ", "Business Intelligence Principal"),
    ("조직: ", "Coupang 3P Marketplace Business Intelligence"),
    ("핵심 역할: ", "3P 셀러 성장 전략 분석, AB Test 설계·실행, BI 대시보드 구축, 프로모션 최적화"),
]
for label, value in coupang_info:
    add_mixed_paragraph(doc, [
        (label, True, Pt(10), BLUE_ACCENT),
        (value, False, Pt(10), DARK_GRAY),
    ], space_before=Pt(2), space_after=Pt(2))

add_shaded_box(doc, "핵심 맥락:",
    "Coupang 3P Marketplace는 수만 셀러의 입점·성장·활성화를 관리하는 플랫폼 비즈니스입니다. "
    "BI Principal로서 '셀러가 어떻게 성장하고, 어디서 이탈하며, 무엇이 매출을 견인하는가'를 "
    "데이터로 증명하고, 실험(AB Test)으로 검증하여, 전략을 실행으로 전환하는 역할을 수행했습니다.")

add_company_highlights(doc, [
    ("FEP", "CPC→FEP 전환 전략으로 17만+ 상품 최적화 — Cohort/AARRR/Funnel/AB Test 종합 적용"),
    ("AARRR", "셀러 온보딩 퍼널을 End-to-End로 측정, 이탈 최다 단계를 식별하여 구조적 UX 개선 제안"),
    ("MyShop", "셀러 브랜드 스토어 런칭 시 AOV 2.4배 상승을 데이터로 검증, 묶음 구매 UX AB Test 설계"),
    ("Auto Price", "자동 가격 조정 서비스의 효과 검증 및 셀러 평가 스코어 로직 구축"),
])

add_thin_gray_line(doc)

# Coupang Project 1: FEP
add_project_section(doc, "Project 1: FEP (Free Exposure Promotion) 상품 최적화", {
    "배경": ("FEP는 셀러가 경쟁력 있는 가격(CPI 90 미만)을 제시하면 추가 노출을 제공하는 3P Marketplace의 "
             "핵심 프로모션 서비스입니다. 판매 시에만 과금되어 셀러 ROAS가 높고, 3P Sales 조직의 핵심 성장 엔진이었습니다."),
    "과정": [
        "Cohort, AARRR, Funnel 방법론을 종합 적용하여 FEP 서비스의 전환·잔존·매출을 다차원 분석",
        "SQL + Tableau 실시간 대시보드 구축, XPC 플랫폼을 통한 AB Test 설계·실행",
        "프로젝트 2년차, 성장 정체 현상을 감지하고 원인 심층 진단에 착수",
    ],
    "문제 발견 및 해결": [
        "CPC에 비용을 투입하면서도 실적이 저조한 상품군(Unit Sold 5개 이하, CPI 100%+)을 세그먼트 분석으로 식별",
        "FEP의 '판매 시 과금' 구조가 CPC 대비 ROAS에서 구조적 우위가 있음을 데이터로 증명",
        "해당 17만여 개 상품을 CPC→FEP로 전환하는 로드맵을 수립, 3P Sales팀과 공동 실행",
    ],
    "사용 도구": "XPC (AB Test Platform), Tableau, SQL (Hive)",
})

add_project_image(doc, os.path.join(IMG_DIR, "coupang_ab_test.png"),
                  "[그림] AB Test Experiment Platform — FEP 프로모션 최적화 실험 결과")
add_project_image(doc, os.path.join(IMG_DIR, "coupang_cpi_dashboard.png"),
                  "[그림] CPI (Coupang Price Index) 가격 분석 대시보드")

# Coupang Project 2: AARRR
add_project_section(doc, "Project 2: 셀러 AARRR 파이프라인 구축", {
    "배경": ("3P Marketplace의 성장은 '좋은 셀러를 유입시키고, 빠르게 활성화시키며, 오래 머물게 하는 것'에 달려 있습니다. "
             "그러나 셀러 라이프사이클(SignUp → Live → Active → GMV) 전체를 채널별로 측정하는 체계가 없었습니다."),
    "핵심 목표 지표": "14주 내 150만 GMV 달성 신규 셀러 수 (Cohort 기반 추적)",
    "문제": "외부 GA 데이터와 내부 DB 간 연결 파이프라인 부재 → 유입 채널별 전환 퍼널 파악 자체가 불가능",
    "해결 과정": [
        "GA 외부 데이터 → 내부 트래커 연결 → Airflow 기반 ETL 파이프라인을 직접 설계·구축",
        "채널별 전환율을 실시간 모니터링하는 Tableau 대시보드 구축, 주간 리포팅 자동화",
    ],
    "핵심 성과": [
        "퍼널 내 이탈 최다 단계가 '사업자 인증'임을 데이터로 식별 → 상품등록 직전으로 순서 변경 제안 → 접근성 개선",
        "14주 150만 GMV 달성 Cohort의 1년 리텐션율이 전체 중 최고임을 검증 — '초기 활성화 속도'가 장기 잔존의 핵심 동인임을 증명",
        "AM 대상 셀러별 맞춤 컨설팅 대시보드 제공 → 셀러 관리를 경험 기반에서 데이터 기반으로 전환",
    ],
    "사용 도구": "Google Analytics, Airflow, Tableau, SQL (Hive)",
})

add_project_image(doc, os.path.join(IMG_DIR, "coupang_seller_funnel.png"),
                  "[그림] 셀러 AARRR 퍼널 — 채널별 SignUp → Live → Active 전환 추적")
add_project_image(doc, os.path.join(IMG_DIR, "coupang_seller_cohort.png"),
                  "[그림] 셀러 Cohort 분석 — SignUp-to-Live 전환율 및 잔존율 추적 대시보드")

# Coupang Project 3: MyShop
add_project_section(doc, "Project 3: MyShop 런칭 프로젝트", {
    "배경": "셀러 자체 브랜드 스토어(MyShop)를 통해 브랜딩 강화와 객단가 상승을 동시에 달성하고자 한 전략적 신규 서비스.",
    "핵심 분석 결과": [
        "MyShop 내 구매 시 AOV(평균 주문단가)가 일반 대비 2.4배 상승한다는 사실을 데이터로 검증 — 서비스 확대의 핵심 근거",
        "합배송 효과와 셀러 쿠폰이 객단가에 미치는 영향을 분리하여 분석, 묶음 구매 UX에 대한 AB Test 설계·실행",
        "AB Test 결과를 기반으로 효과가 큰 셀러 세그먼트를 식별하고 단계적 확산 전략 수립",
    ],
    "사용 도구": "XPC (AB Test Platform), Tableau, SQL (Hive)",
})

add_project_image(doc, os.path.join(IMG_DIR, "coupang_gmv_dashboard.png"),
                  "[그림] GMV per AM 성과 대시보드 — 셀러별 매출 기여도 분석")

# Coupang Project 4: Auto Price
add_project_section(doc, "Project 4: Auto Price Service 분석 및 셀러 스코어 관리", {
    "과정": [
        "Auto Price 이용 셀러 vs 비이용 셀러 간 매출·전환율·가격 변동 패턴을 비교 분석하고 AB Test를 설계하여 서비스 효과를 정량 검증",
        "매출, 별점, 배송 타입, EDD/PDD 등 다차원 요소를 반영한 셀러 종합 평가 스코어 산출 로직 설계·구현",
        "정책 위반 셀러 자동 탐지 및 단계별 제재(경고 → 노출 제한 → 블록) 프로세스 설계",
    ],
    "사용 도구": "SQL (Hive), Tableau, XPC (AB Test Platform)",
})

add_project_image(doc, os.path.join(IMG_DIR, "coupang_multi_dashboard.png"),
                  "[그림] Coupang BI 종합 분석 대시보드 — 다차원 KPI 모니터링")


# ════════════════════════════════════════════════════════════════
#  PAGE BREAK → EBAY KOREA
# ════════════════════════════════════════════════════════════════
doc.add_page_break()

doc.add_heading("3. eBay Korea (2014.07 ~ 2020.03)", level=2)

add_styled_paragraph(doc, "Gmarket, Auction, G9 — 대한민국 대표 e-Commerce 3개 사이트 운영. 매출 1조 원, 직원 약 1,000명 (2018년 기준).",
                    size=Pt(10), color=MEDIUM_GRAY, space_before=Pt(4), space_after=Pt(8))

ebay_info = [
    ("직급/역할: ", "과장 (Data PM)"),
    ("조직: ", "Data & Martech"),
    ("핵심 역할: ", "행동데이터 수집 시스템 구축(Owner PM), 데이터 거버넌스 표준화, BI 플랫폼 구축, Fraud Detection 기획"),
]
for label, value in ebay_info:
    add_mixed_paragraph(doc, [
        (label, True, Pt(10), BLUE_ACCENT),
        (value, False, Pt(10), DARK_GRAY),
    ], space_before=Pt(2), space_after=Pt(2))

add_shaded_box(doc, "핵심 맥락:",
    "eBay Korea는 3개 사이트가 각각 다른 외부 솔루션으로 데이터를 수집하고, 부서마다 지표 정의가 달랐습니다. "
    "'전사적으로 비교 가능한 데이터 기반'이 없었던 것입니다. 이 환경에서 데이터 수집 시스템을 직접 기획하고, "
    "거버넌스를 통일하며, BI 플랫폼을 구축하는 — 데이터 인프라를 0에서 1로 만드는 경험을 했습니다.")

add_company_highlights(doc, [
    ("Montelena", "3개 사이트 행동데이터 수집·분석 시스템 기획 (Owner PM, 1.5년) — 데이터 거버넌스 통일의 기반"),
    ("Deep Link", "Web↔App 데이터 추적 단절 문제를 구조적으로 해결 → 2019 2Q PIC Unit Award 수상"),
    ("BI Platform", "검색 키워드·프로모션 실적 분석을 위한 셀프서비스 BI 플랫폼 구축"),
    ("Gotham", "Fraud Detection 관리 도구 'Gotham' 기획 — ATO 이상 행동 탐지 로직 설계"),
    ("Data TF", "전사 Data System 3개년 개발 로드맵 수립에 핵심 기여"),
])

add_thin_gray_line(doc)

# Phase I
doc.add_heading("Phase I: Data Insight POD / Data PM (2016.07 ~ 2020.03)", level=3)

add_project_section(doc, "Project 1: Montelena TF — 행동데이터 수집·분석 시스템 구축 (Owner PM, 1.5년)", {
    "배경": ("Gmarket, Auction, G9가 각각 다른 솔루션으로 데이터를 수집하고, 부서마다 PV, UV 등의 정의가 달라 "
             "전사 차원의 데이터 비교 자체가 불가능한 상황이었습니다."),
    "목표": "3개 사이트의 행동데이터 수집 통일 + 지표 표준화(거버넌스) + 분석 서비스 제공까지 End-to-End 구현",
    "해결 과정": [
        "[데이터 거버넌스] PGUID↔CUID 연결 체계 설계, PV/UV/Click/Members 등 8개 핵심 지표의 Standard Query 정립 — 이 경험이 이후 Tving AI 쿼리 시스템 설계의 기반",
        "[태깅 시스템] 웹 컨테이너 자동 태깅을 시도했으나 레거시 시스템 제약으로 수동 태깅으로 전환, 비즈니스 임팩트가 큰 16개 핵심 도메인에 집중",
        "[도메인 밸류] 행동데이터와 트랜잭션 데이터를 연결하여, 각 도메인(검색, 카테고리, 프로모션 등)의 구매 기여도를 정량적으로 산출",
        "[전파] 전사 PM Unit 대상 Tagging System 프레젠테이션 진행 — 조직 전체의 데이터 활용도 향상에 기여",
    ],
    "사용 도구": "Hive SQL, Tableau, ElasticSearch, Python (Py-Hive)",
})

add_project_image(doc, os.path.join(IMG_DIR, "ebay_dashboard_design.png"),
                  "[그림] 데이터 서비스 플랫폼 설계 — 대시보드 UI/UX 기획 문서")

add_project_section(doc, "Project 2: Deep Link Gate (Main PM)", {
    "배경": "웹에서 앱으로 이동할 때 데이터 추적이 끊어져, 마케팅 채널별 실제 ROI를 측정할 수 없는 문제",
    "과정": [
        "Web→App 전환 시 사용자 식별이 유지되는 딥링크 게이트 로직을 설계·개발",
        "유입 케이스별(직접 접속, 검색, 광고, 푸시) AB Test를 설계·실행하여 최적 전환 경로를 검증",
        "전사 대상 MKT Deep Link Logic 프레젠테이션 진행",
    ],
    "성과": "2019년 2분기 PIC Unit Award 수상 — Web↔App 간 데이터 추적 체계를 구조적으로 완성한 공로",
    "사용 도구": "Hadoop, Hive SQL, Tableau",
})

add_project_section(doc, "Project 3: BI 서비스 플랫폼 구축", {
    "구축 내용": [
        "[검색 키워드 분석] 일간/주간/월간 키워드별 검색 정보 셀프서비스 플랫폼 (Tableau, HIVE, Bootstrap)",
        "[프로모션 실적 분석] 방문당(Per Visit) 프로모션 실적 다차원 분석 플랫폼",
    ],
})

add_project_image(doc, os.path.join(IMG_DIR, "ebay_tableau_kpi.png"),
                  "[그림] Tableau KPI Summary 대시보드")
add_project_image(doc, os.path.join(IMG_DIR, "ebay_tableau_trends.png"),
                  "[그림] Tableau Trend Analysis 대시보드")

add_project_section(doc, "Project 4: Fraud Detection \"Gotham\" 기획", {
    "과정": [
        "정상 사용자 행동 패턴을 프로파일링하고, ATO(Account Takeover) 판별 기준을 데이터 기반으로 수립",
        "IP, Cookie, 접속 시간 패턴(Raw Duration) 등 복합 시그널 기반 이상 탐지 로직 개발",
        "Fraud Detection 관리 도구 'Gotham' 기획 — 실시간 알림, 사례별 관리 워크플로우, 대응 자동화",
    ],
    "사용 도구": "ElasticSearch, Kibana, SQL, Python",
})

add_project_image(doc, os.path.join(IMG_DIR, "ebay_gotham.png"),
                  "[그림] Fraud Detection 'Gotham' — 이상 거래 탐지 UI")

add_project_section(doc, "Project 5: 다양한 데이터 분석 및 리포팅", {
    "수행 내용": [
        "Funnel 기반 Domain별 이탈/전환 분석 (Oracle SQL, Hive, Zeplin)",
        "Global KPI 'Searched Session Per GMV' 기획 (MySQL, Hive, Tableau, Zeplin)",
        "Domain Value 지표 기획 — 각 도메인의 매출 기여도 정량 산출 (Hive SQL, Python)",
        "Session 정합성 관제 리포팅 (ElasticSearch, Kibana, D3.js)",
    ],
})

add_thin_gray_line(doc)

# Phase II
doc.add_heading("Phase II: 전사 Data TF (2016.01 ~ 2016.07)", level=3)

add_project_section(doc, "전사 Data System 3개년 개발계획 수립", {
    "수행 내용": [
        "Business Unit별 데이터 요구사항 조사 (마케팅, 영업, 상품, CS 등 전 부서)",
        "Data System Module 요구조건 기획 및 개발 우선순위 확정",
        "전사 Data System 3개년 로드맵 수립에 기여",
    ],
})

add_thin_gray_line(doc)

# Phase III
doc.add_heading("Phase III: Sales 전략기획팀 (2014.07 ~ 2016.01)", level=3)

add_project_section(doc, "Sales 전략기획 및 데이터 분석", {
    "수행 내용": [
        "상품 Category 재분류를 위한 연관성 분석 — Association Rule 기반 (R, MySQL)",
        "Super Deal 구좌 분배 최적화 — 노출 대비 전환율을 최대화하는 배분 로직 설계 (Hive SQL, MySQL, Python)",
        "Membership MKT '취미 연구소' 기획·실행, O2O 스마트 모빌리티 프로젝트 기획/실행",
        "52주 Seasonal 패턴 분석, 쿠폰 스펙(할인율, 최소 결제금액)과 타겟 그룹별 최적 조합 도출",
    ],
    "핵심 성과": "2015년 eBay Korea 영업BU 최우수 직원상 수상 — 데이터 기반 Sales 전략 기여를 인정",
})

add_project_image(doc, os.path.join(IMG_DIR, "ebay_r_analysis.png"),
                  "[그림] R 기반 상품 Category 연관성 분석 — Association Rule 시각화")


# ════════════════════════════════════════════════════════════════
#  PAGE BREAK → E-LAND
# ════════════════════════════════════════════════════════════════
doc.add_page_break()

doc.add_heading("4. E-Land Retail (2010.11 ~ 2014.07)", level=2)

add_styled_paragraph(doc, "NC백화점, 2001아울렛, New-Wave 등 보유. 매출 2조 1천억 원, 직원 약 4,000명.",
                    size=Pt(10), color=MEDIUM_GRAY, space_before=Pt(4), space_after=Pt(8))

eland_info = [
    ("직급/역할: ", "주임"),
    ("조직: ", "Marketing / CRM"),
    ("핵심 역할: ", "신규 점포 오픈 마케팅, 고객 세그먼트 기반 CRM, 전국 재고 순환 시스템 구축"),
]
for label, value in eland_info:
    add_mixed_paragraph(doc, [
        (label, True, Pt(10), BLUE_ACCENT),
        (value, False, Pt(10), DARK_GRAY),
    ], space_before=Pt(2), space_after=Pt(2))

add_shaded_box(doc, "핵심 맥락:",
    "커리어의 출발점에서 '고객 데이터가 실제 매출에 어떻게 연결되는가'를 오프라인 현장에서 체득한 시기입니다. "
    "상권 분석, CRM 세그먼트 전략, 재고 최적화 등 데이터 분석의 기본기를 실행 중심으로 다졌습니다.")

add_company_highlights(doc, [
    ("Open MKT", "NC New-Wave 신규 점포 오픈 마케팅 — 상권 분석, 예상 매출 산출, 고객 타겟팅 CRM 기획"),
    ("Inventory", "전국 5개 창고·48개 점포, 800억 원 규모 의류 재고의 최적 순환 시스템 구축"),
])

add_thin_gray_line(doc)

doc.add_heading("Phase I: NC New-Wave 점포 런칭/CRM 기획 (2013.01 ~ 2014.07)", level=3)

add_project_section(doc, "신규 점포 런칭 마케팅 및 CRM 기획", {
    "수행 내용": [
        "신규 출점 기획 전 과정 수행 — 상권 분석, 경쟁사 현황 조사, 예상 매출 산출, 마케팅 예산 편성·관리",
        "주소지 기반 고객 특성 분석으로 지역별 타겟 마케팅 전략 수립 — 이후 eBay, Coupang에서의 세그먼트 분석의 출발점",
        "Membership CRM 기획 — 시즌별 DM 기획, 층별 Main Event 기획 및 실행",
        "고객 세그먼트별 맞춤형 프로모션 설계 및 효과 측정(반응률, 객단가 변화)",
    ],
})

add_thin_gray_line(doc)

doc.add_heading("Phase II: 의류 재고 순환 System TF (2012.01 ~ 2013.01)", level=3)

add_project_section(doc, "전사 의류 재고 순환 최적화 시스템 구축", {
    "배경": "전국 5개 물류 창고에 약 800억 원 규모 의류 재고가 분산되어 있었으나, 48개 점포별 최적 배분 기준이 없는 상태",
    "수행 내용": [
        "5개 창고 재고 현황과 48개 점포별 판매 데이터를 분석하여 최적 이동 비용·루트 설계",
        "연간 재고 순환 스케줄 수립 — 물류 비용 절감과 판매 기회 극대화를 동시에 달성하는 균형점 도출",
        "시즌별 단계적 할인 정책 기획 — 재고 소진율을 높이면서도 마진을 최대한 보전하는 구간별 전략 설계",
    ],
})


# ════════════════════════════════════════════════════════════════
#  PAGE BREAK → AWARDS & TIMELINE & CLOSING
# ════════════════════════════════════════════════════════════════
doc.add_page_break()

# Awards
doc.add_heading("수상 이력", level=1)
add_blue_line(doc)

awards_table = doc.add_table(rows=3, cols=4)
awards_table.alignment = WD_TABLE_ALIGNMENT.CENTER
format_table_borders(awards_table, TABLE_BORDER_COLOR)
awards_headers = ["연도", "수상명", "수상 배경", "소속"]
awards_data = [
    ["2019", "2Q PIC Unit Award", "Deep Link Gate — 웹→앱 데이터 추적 단절 문제 해결", "eBay Korea"],
    ["2015", "영업BU 최우수 직원상", "Membership MKT '취미 연구소' 기획·실행 — 매출 기여도 제고", "eBay Korea"],
]
for j, h in enumerate(awards_headers):
    cell = awards_table.rows[0].cells[j]
    set_cell_shading(cell, TABLE_HEADER_BG)
    format_cell_text(cell, h, bold=True, size=Pt(10), color=WHITE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
awards_table.rows[0].cells[0].width = Cm(2.0)
awards_table.rows[0].cells[1].width = Cm(4.0)
awards_table.rows[0].cells[2].width = Cm(8.0)
awards_table.rows[0].cells[3].width = Cm(2.5)
for i, row_data in enumerate(awards_data, 1):
    for j, cell_text in enumerate(row_data):
        cell = awards_table.rows[i].cells[j]
        if i % 2 == 0: set_cell_shading(cell, TABLE_ALT_BG)
        alignment = WD_ALIGN_PARAGRAPH.CENTER if j in [0, 3] else WD_ALIGN_PARAGRAPH.LEFT
        format_cell_text(cell, cell_text, bold=(j==1), size=Pt(10), color=DARK_GRAY, alignment=alignment)

# Timeline
doc.add_heading("커리어 타임라인", level=1)

timeline_data = [
    ("2010.11", "E-Land Retail 입사", "Marketing / CRM 주임"),
    ("2012.01", "의류 재고 순환 System TF", "800억 규모 재고 최적화"),
    ("2013.01", "NC New-Wave 신규 점포 런칭", "상권 분석 / Open MKT / CRM"),
    ("2014.07", "eBay Korea 입사", "Sales 전략기획팀"),
    ("2015", "영업BU 최우수 직원상 수상", "'취미 연구소' MKT 기획"),
    ("2016.01", "전사 Data TF 참여", "Data System 3개년 계획 수립"),
    ("2016.07", "Data Insight POD 이동", "Data PM으로 역할 전환"),
    ("2017~18", "Montelena TF 수행", "유저 행동데이터 수집/분석 시스템 (1.5년)"),
    ("2019", "Deep Link Gate / PIC Award", "웹-앱 데이터 추적 체계 완성"),
    ("2020.03", "Coupang 입사", "BI Principal"),
    ("2020~25", "3P Marketplace BI 리드", "FEP, AARRR, MyShop, Auto Price"),
    ("2025.07", "Tving 입사", "Data Analyst Team 팀장"),
    ("현재", "AI 분석 체계 구축 리딩", "AI 대시보드, 예측 모델링, KPI OKR, 팀 매니징"),
]

tl_table = doc.add_table(rows=len(timeline_data) + 1, cols=3)
tl_table.alignment = WD_TABLE_ALIGNMENT.CENTER
format_table_borders(tl_table, TABLE_BORDER_COLOR)
for j, h in enumerate(["시기", "이벤트", "핵심 내용"]):
    cell = tl_table.rows[0].cells[j]
    set_cell_shading(cell, TABLE_HEADER_BG)
    format_cell_text(cell, h, bold=True, size=Pt(9.5), color=WHITE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
tl_table.rows[0].cells[0].width = Cm(2.5)
tl_table.rows[0].cells[1].width = Cm(5.5)
tl_table.rows[0].cells[2].width = Cm(8.5)
for i, (period, event, detail) in enumerate(timeline_data, 1):
    cells = tl_table.rows[i].cells
    if i % 2 == 0:
        for c in cells: set_cell_shading(c, TABLE_ALT_BG)
    format_cell_text(cells[0], period, bold=True, size=Pt(9), color=BLUE_ACCENT, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    format_cell_text(cells[1], event, bold=True, size=Pt(9), color=DARK_GRAY, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    format_cell_text(cells[2], detail, bold=False, size=Pt(9), color=MEDIUM_GRAY, alignment=WD_ALIGN_PARAGRAPH.LEFT)

# Closing
add_thin_gray_line(doc)

closing_text = (
    "오프라인 리테일에서 고객과 재고를 분석하며 데이터의 기본기를 다졌고, "
    "e-Commerce에서 수집 시스템과 거버넌스를 0에서 구축하며 데이터 인프라의 설계자가 되었고, "
    "Coupang에서 실험 기반의 전략적 BI를 리드했으며, "
    "지금은 Tving에서 AI가 스스로 분석하고 답하는 체계를 만들고 있습니다. "
    "이 여정에서 일관된 것은, 데이터를 '보는 것'에서 '작동하는 것'으로 바꾸는 실행력입니다."
)
add_styled_paragraph(doc, closing_text, size=Pt(10), color=MEDIUM_GRAY,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(20),
                    space_after=Pt(8), line_spacing=1.5)

add_styled_paragraph(doc, "2026년 2월", size=Pt(10.5), color=DARK_GRAY,
                    alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_before=Pt(20), space_after=Pt(4))
add_styled_paragraph(doc, "연새한솔", size=Pt(14), bold=True, color=DARK_GRAY,
                    alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_before=Pt(0), space_after=Pt(0))


# ════════════════════════════════════════════════════════════════
#  SAVE
# ════════════════════════════════════════════════════════════════
output_path = r"C:/Users/user/Desktop/이력서/202502_연새한솔_경력기술서_Portfolio.docx"
doc.save(output_path)
print(f"Document saved successfully: {output_path}")
